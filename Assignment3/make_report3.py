#!/usr/bin/env python3
import os
from docx import Document
from docx.shared import Pt, RGBColor, Inches, Cm
from docx.enum.text import WD_ALIGN_PARAGRAPH

OUT = "/home/nouman-ejaz/Desktop/Information Security Assigment 1/Assignment3"

doc = Document()
for section in doc.sections:
    section.top_margin    = Cm(2.0)
    section.bottom_margin = Cm(2.0)
    section.left_margin   = Cm(2.5)
    section.right_margin  = Cm(2.5)

def heading(text, level=1, color=None):
    p = doc.add_heading(text, level=level)
    p.alignment = WD_ALIGN_PARAGRAPH.LEFT
    if color:
        for run in p.runs:
            run.font.color.rgb = RGBColor(*color)
    return p

def body(text, bold=False):
    p = doc.add_paragraph()
    run = p.add_run(text)
    run.font.size = Pt(11)
    run.font.bold = bold
    return p

def bullet(text):
    p = doc.add_paragraph(text, style="List Bullet")
    p.runs[0].font.size = Pt(11)
    return p

# ── Title Page ─────────────────────────────────────────────────────
doc.add_paragraph(); doc.add_paragraph()
for text, sz, bold in [
    ("Assignment 3", 26, True),
    ("DefendX: Network Security Scanner & Firewall Visualizer", 18, True),
]:
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    r = p.add_run(text)
    r.font.size = Pt(sz); r.font.bold = bold
    r.font.color.rgb = RGBColor(0x1A,0x23,0x7E)

doc.add_paragraph()
for label, val in [
    ("Student Name:", "Abdul Rehman"),
    ("Roll Number:", "03-134222-005"),
    ("Course:", "Information Security (CLO 4)"),
    ("Submission:", "Single-page Python Flask tool with report"),
]:
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    rb = p.add_run(f"{label} "); rb.font.bold = True; rb.font.size = Pt(12)
    rv = p.add_run(val); rv.font.size = Pt(12)

doc.add_page_break()

# ── 1. Introduction ────────────────────────────────────────────────
heading("1. Introduction", 1, (0x1A,0x23,0x7E))
body(
    "DefendX is a comprehensive Python Flask-based desktop and web security utility that provides network "
    "scanning and firewall simulation features. Built to satisfy the requirements of CLO 4, the application "
    "empowers users to scan specific targets (localhost or network addresses) to detect available ports and "
    "services. It also includes an advanced visualizer that tests policy rules on incoming and outgoing "
    "traffic packets."
)

# ── 2. Scanning & Testing Techniques ───────────────────────────────
doc.add_paragraph()
heading("2. Network Scanning & Evaluation", 1, (0x1A,0x23,0x7E))

heading("2.1 TCP Connect Scanning", 2, (0x0D,0x47,0xA1))
body(
    "The core scanner utilizes the full three-way TCP handshake (SYN, SYN-ACK, ACK) to reliably establish "
    "connections. If the socket connects successfully, the port is identified as 'Open'. If an error code is "
    "returned or a timeout occurs, the port status defaults to 'Closed'. This is completely non-privileged "
    "and runs natively across any OS without root/sudo access."
)

heading("2.2 Firewall Simulation Logic", 2, (0x0D,0x47,0xA1))
body(
    "The firewall rule logic chains together priority-based rules using a top-down evaluation strategy. "
    "Each rule can explicitly define an allow or deny policy targeting specific ports or entire source IP "
    "ranges. The simulation validates whether traffic passes or gets intercepted by matching parameters "
    "sequentially against defined rules."
)

# ── 3. Application Implementation ──────────────────────────────────
doc.add_page_break()
heading("3. Technical Implementation & Flow", 1, (0x1A,0x23,0x7E))

heading("3.1 Technology Stack", 2, (0x0D,0x47,0xA1))
body("Backend: Python 3.x, Flask (REST APIs for scanning)")
body("Frontend: Responsive HTML5, Vanilla CSS3 for glassmorphism layout, JavaScript ES2022")
body("Libraries: standard Python socket module")

heading("3.2 Core Features", 2, (0x0D,0x47,0xA1))
for feat in [
    "Integrated target scanner with single-page frontend interface",
    "Live visual feedback using connection line coloring (Red for blocked, Green for allowed)",
    "Custom rule insertion: allow or deny access based on precise IP or port",
    "Comprehensive tabular display of target, service name, and scan type",
    "Dynamic traffic interception simulation when scanning live targets",
]:
    bullet(feat)

# ── 4. Conclusion ──────────────────────────────────────────────────
doc.add_paragraph()
heading("4. Conclusion", 1, (0x1A,0x23,0x7E))
body(
    "DefendX successfully combines direct network scanning with dynamic firewall rule visual testing "
    "to provide a fully self-contained security simulator. The tool accurately illustrates rule chains "
    "and prioritizations, helping students understand defensive network design principles."
)

path = f"{OUT}/Assignment3_Report_AbdulRehman.docx"
doc.save(path)
print(f"✅ Word report saved: {path}")
