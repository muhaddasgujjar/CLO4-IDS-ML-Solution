#!/usr/bin/env python3
"""Generate the Assignment 2 Word report."""
import os
from docx import Document
from docx.shared import Pt, RGBColor, Inches, Cm
from docx.enum.text import WD_ALIGN_PARAGRAPH

OUT = "/home/nouman-ejaz/Desktop/Information Security Assigment 1/Assignment2"

doc = Document()
for section in doc.sections:
    section.top_margin    = Cm(2.0)
    section.bottom_margin = Cm(2.0)
    section.left_margin   = Cm(2.5)
    section.right_margin  = Cm(2.5)

def heading(text, level=1, color=None):
    p = doc.add_heading(text, level=level)
    p.alignment = WD_ALIGN_PARAGRAPH.LEFT
    if color:
        for run in p.runs:
            run.font.color.rgb = RGBColor(*color)
    return p

def body(text, bold=False):
    p = doc.add_paragraph()
    run = p.add_run(text)
    run.font.size = Pt(11)
    run.font.bold = bold
    return p

def bullet(text):
    p = doc.add_paragraph(text, style="List Bullet")
    p.runs[0].font.size = Pt(11)
    return p

def add_table(headers, rows_data):
    from docx.enum.table import WD_TABLE_ALIGNMENT
    table = doc.add_table(rows=1, cols=len(headers))
    table.style = "Table Grid"
    table.alignment = WD_TABLE_ALIGNMENT.CENTER
    hdr = table.rows[0].cells
    for i, h in enumerate(headers):
        hdr[i].text = h
        hdr[i].paragraphs[0].runs[0].font.bold = True
        hdr[i].paragraphs[0].runs[0].font.size = Pt(10)
        hdr[i].paragraphs[0].alignment = WD_ALIGN_PARAGRAPH.CENTER
    for row in rows_data:
        cells = table.add_row().cells
        for i, val in enumerate(row):
            cells[i].text = str(val)
            cells[i].paragraphs[0].runs[0].font.size = Pt(10)
            cells[i].paragraphs[0].alignment = WD_ALIGN_PARAGRAPH.CENTER
    return table

# ── Title Page ─────────────────────────────────────────────────────
doc.add_paragraph(); doc.add_paragraph()
for text, sz, bold in [
    ("Assignment 2", 26, True),
    ("CryptoVault: Web-Based Text Encryption Tool", 18, True),
]:
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    r = p.add_run(text)
    r.font.size = Pt(sz); r.font.bold = bold
    r.font.color.rgb = RGBColor(0x1A,0x23,0x7E)

doc.add_paragraph()
for label, val in [
    ("Student Name:", "Abdul Rehman"),
    ("Roll Number:", "03-134222-005"),
    ("Course:", "Information Security"),
    ("Submission:", "Single-page web application with report"),
]:
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    rb = p.add_run(f"{label} "); rb.font.bold = True; rb.font.size = Pt(12)
    rv = p.add_run(val); rv.font.size = Pt(12)

doc.add_page_break()

# ── 1. Introduction ────────────────────────────────────────────────
heading("1. Introduction", 1, (0x1A,0x23,0x7E))
body(
    "CryptoVault is a fully client-side, single-page web application that allows users to encrypt, decrypt, "
    "and hash text using six cryptographic algorithms. Developed as part of the Information Security course, "
    "it demonstrates the practical application of both classical and modern cryptographic techniques. "
    "All operations execute entirely within the user's browser using JavaScript — no data is ever transmitted "
    "to any server, ensuring complete privacy."
)

# ── 2. Encryption Techniques ───────────────────────────────────────
doc.add_paragraph()
heading("2. Encryption Techniques Used", 1, (0x1A,0x23,0x7E))

# Caesar
heading("2.1 Caesar Cipher", 2, (0x0D,0x47,0xA1))
body("Type: Classical Symmetric Substitution Cipher", bold=True)
body(
    "The Caesar Cipher is one of the earliest known encryption techniques, attributed to Julius Caesar. "
    "It shifts each alphabetic character by a fixed integer value (shift key, 1–25) in the alphabet. "
    "Non-alphabetic characters (digits, punctuation, spaces) are left unchanged."
)
body("Encryption formula: E(x) = (x + shift) mod 26", bold=False)
body("Decryption formula: D(x) = (x − shift + 26) mod 26", bold=False)
body(
    "Security Note: Caesar Cipher has a key space of only 25 possible shifts, making it trivially breakable "
    "by brute-force or frequency analysis. It is included for educational and historical context."
)

# AES
doc.add_paragraph()
heading("2.2 AES-256 (Advanced Encryption Standard)", 2, (0x0D,0x47,0xA1))
body("Type: Modern Symmetric Block Cipher | Mode: CBC | Key Size: 256-bit", bold=True)
body(
    "AES is the NIST-standardized symmetric encryption algorithm (FIPS 197, 2001) and the global "
    "gold standard for data-at-rest and data-in-transit encryption. It uses a 256-bit key and operates "
    "on 128-bit blocks. CryptoVault uses CBC (Cipher Block Chaining) mode with a random 128-bit IV "
    "(Initialization Vector) prepended to the ciphertext to prevent identical plaintexts from producing "
    "identical ciphertexts. The IV is encoded as a hex prefix separated by ':' for storage and recovery."
)
body("Key derivation: The user key is padded/truncated to exactly 32 bytes (256 bits).")
body("Real-world usage: TLS 1.3, WPA3, disk encryption (BitLocker, VeraCrypt), banking systems.")

# DES
doc.add_paragraph()
heading("2.3 DES (Data Encryption Standard)", 2, (0x0D,0x47,0xA1))
body("Type: Legacy Symmetric Block Cipher | Mode: CBC | Key Size: 56-bit", bold=True)
body(
    "DES was the US federal standard from 1977 until replaced by AES in 2001. It uses a 56-bit key "
    "and 64-bit blocks. DES is no longer considered secure — its 56-bit key was first publicly broken "
    "in 22 hours by EFF's Deep Crack in 1998. CryptoVault implements DES CBC with a random 8-byte IV "
    "for demonstration and historical comparison. In practice, Triple-DES (3DES) or AES should be used."
)
body("Security Note: Do not use DES for any production data. Included for educational purposes only.")

# RSA
doc.add_paragraph()
heading("2.4 RSA (Rivest–Shamir–Adleman)", 2, (0x0D,0x47,0xA1))
body("Type: Asymmetric Public-Key Cipher | Padding: OAEP-SHA256 | Key Size: 512/1024-bit", bold=True)
body(
    "RSA is the most widely deployed asymmetric (public-key) cryptosystem, invented in 1977. It is based "
    "on the mathematical difficulty of factoring the product of two large prime numbers. CryptoVault generates "
    "an ephemeral RSA key pair using the browser's built-in Web Crypto API (SubtleCrypto.generateKey). "
    "The public key encrypts the message; the private key decrypts it. OAEP (Optimal Asymmetric Encryption "
    "Padding) with SHA-256 is used to prevent padding oracle attacks."
)
body(
    "The private key is stored in sessionStorage for the duration of the browser session, allowing decryption "
    "in the same tab. Note: RSA is typically used only to encrypt small payloads (e.g., symmetric session keys) "
    "in hybrid encryption schemes such as TLS."
)

# Base64
doc.add_paragraph()
heading("2.5 Base64 Encoding", 2, (0x0D,0x47,0xA1))
body("Type: Binary-to-Text Encoding (NOT encryption)", bold=True)
body(
    "Base64 is a group of binary-to-text encoding schemes that represent binary data using 64 printable ASCII "
    "characters (A–Z, a–z, 0–9, +, /). It is NOT encryption — it provides no confidentiality. Base64 is "
    "fully reversible without any key. It is used extensively to safely transmit binary data over "
    "text-based protocols: email attachments (MIME), embedding images in HTML (data URIs), JWT tokens, "
    "and API payloads."
)

# SHA-256
doc.add_paragraph()
heading("2.6 SHA-256 (Secure Hash Algorithm 256)", 2, (0x0D,0x47,0xA1))
body("Type: Cryptographic Hash Function | Output: 256-bit (64 hex chars) | One-Way", bold=True)
body(
    "SHA-256 is part of the SHA-2 family standardized by NIST. Unlike the algorithms above, it is a "
    "one-way function — given a hash, it is computationally infeasible to recover the original input "
    "(preimage resistance). It also provides collision resistance: no two different inputs should produce "
    "the same hash."
)
body("SHA-256 is used for: password storage (with salt), file integrity verification, digital signatures, "
     "Bitcoin/blockchain proof-of-work, and HMAC constructions.")

# ── 3. Implementation ──────────────────────────────────────────────
doc.add_page_break()
heading("3. Application Implementation", 1, (0x1A,0x23,0x7E))

heading("3.1 Tech Stack", 2, (0x0D,0x47,0xA1))
add_table(
    ["Layer","Technology","Purpose"],
    [
        ["Frontend","HTML5 + CSS3","Structure and dark-theme UI"],
        ["Logic","JavaScript (ES2020)","Encryption engine, UI state"],
        ["Crypto Library","CryptoJS 4.2.0","AES, DES, SHA-256, Base64"],
        ["RSA","Web Crypto API (SubtleCrypto)","RSA-OAEP key generation and encryption"],
        ["Fonts","Google Fonts (Inter, JetBrains Mono)","Typography"],
    ]
)

doc.add_paragraph()
heading("3.2 Core Features", 2, (0x0D,0x47,0xA1))
for feat in [
    "Three Operation Modes: Encrypt, Decrypt, Hash — via a tab-based interface",
    "Six Algorithms: Caesar, AES-256, DES, RSA, Base64, SHA-256",
    "Input Validation: Prevents empty submission, enforces algorithm selection and key entry",
    "Auto Key Generator: Generates cryptographically random keys for AES (32 bytes) and DES (8 bytes)",
    "IV Handling: AES and DES automatically prepend the random IV to ciphertext (hex:ciphertext format)",
    "Copy to Clipboard: One-click copy of output with toast notification (bonus feature)",
    "Stats Bar: Displays algorithm, mode, input/output character counts after each operation",
    "Algorithm Info Cards: Educational descriptions of each algorithm with security classifications",
    "Responsive Design: Works on desktop and mobile screens",
    "Zero Server Dependency: Fully offline-capable — open as a local HTML file",
]:
    bullet(feat)

heading("3.3 Validation Logic", 2, (0x0D,0x47,0xA1))
body("The application enforces the following validations before processing:")
for v in [
    "Algorithm must be selected from dropdown (not empty default)",
    "Input text cannot be empty",
    "Secret key required for AES and DES algorithms",
    "Hash mode automatically disabled for SHA-256 in encrypt/decrypt tabs",
    "RSA ciphertext must start with 'RSA:' prefix for decryption",
    "Base64 decode validates input is valid Base64 before processing",
]:
    bullet(v)

# ── 4. Security Analysis ───────────────────────────────────────────
doc.add_paragraph()
heading("4. Security Analysis", 1, (0x1A,0x23,0x7E))
add_table(
    ["Algorithm","Security Level","Key Space","Recommended Use"],
    [
        ["Caesar Cipher","❌ Broken","25 keys","Education only"],
        ["DES","⚠️ Weak","2^56 ≈ 72 quadrillion","Legacy / Education"],
        ["Base64","➖ None (encoding)","N/A","Data transport"],
        ["RSA-1024","⚠️ Borderline","~1024-bit prime product","Short-term only"],
        ["AES-256","✅ Strong","2^256 keys","Production use"],
        ["SHA-256","✅ Strong (hash)","256-bit output","Password hashing + integrity"],
    ]
)

# ── 5. Conclusion ──────────────────────────────────────────────────
doc.add_paragraph()
heading("5. Conclusion", 1, (0x1A,0x23,0x7E))
body(
    "CryptoVault successfully implements all core and bonus requirements of Assignment 2. The application "
    "demonstrates a spectrum of cryptographic techniques — from the historically significant Caesar Cipher "
    "to the NIST-standard AES-256 — within an intuitive, modern web interface. By implementing all logic "
    "client-side using CryptoJS and the Web Crypto API, the tool ensures user data never leaves the browser. "
    "The inclusion of decryption, SHA-256 hashing, and one-click clipboard copy fulfills all bonus criteria. "
    "This project solidifies the understanding of how real-world applications protect data confidentiality "
    "and integrity through cryptographic primitives."
)

# ── References ─────────────────────────────────────────────────────
doc.add_paragraph()
heading("References", 1, (0x1A,0x23,0x7E))
refs = [
    "NIST FIPS 197 – Advanced Encryption Standard (AES), 2001. https://csrc.nist.gov/publications/fips/fips197/fips-197.pdf",
    "NIST FIPS 46-3 – Data Encryption Standard (DES), 1999.",
    "Rivest, R. L., Shamir, A., & Adleman, L. (1978). A method for obtaining digital signatures and public-key cryptosystems. CACM, 21(2), 120–126.",
    "CryptoJS Library Documentation. https://cryptojs.gitbook.io/",
    "MDN Web Docs – SubtleCrypto API. https://developer.mozilla.org/en-US/docs/Web/API/SubtleCrypto",
    "NIST FIPS 180-4 – Secure Hash Standard (SHS / SHA-256), 2015.",
]
for i, ref in enumerate(refs, 1):
    p = doc.add_paragraph(f"[{i}] {ref}")
    p.runs[0].font.size = Pt(10)

doc.add_paragraph()
p = doc.add_paragraph()
rb = p.add_run("GitHub Repository: "); rb.font.bold = True
rv = p.add_run("https://github.com/muhaddasgujjar/CLO4-IDS-ML-Solution/tree/main/Assignment2")
rv.font.color.rgb = RGBColor(0x19,0x76,0xD2)

path = f"{OUT}/Assignment2_Report_AbdulRehman.docx"
doc.save(path)
print(f"✅ Word report saved: {path}")
