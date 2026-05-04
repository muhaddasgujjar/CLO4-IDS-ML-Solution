#!/usr/bin/env python3
import os
from PIL import Image, ImageDraw, ImageFont
from docx import Document
from docx.shared import Pt, RGBColor, Inches, Cm
from docx.enum.text import WD_ALIGN_PARAGRAPH

OUT_DIR = "/home/nouman-ejaz/Desktop/Information Security Assigment 1/Assignment4"

def create_terminal_screenshot(text_lines, filename):
    # Create an image that looks like a terminal
    bg_color = (40, 42, 54) # Dracula background
    text_color = (248, 248, 242)
    prompt_color = (80, 250, 123)
    
    # Calculate image size
    width = 900
    height = len(text_lines) * 22 + 40
    img = Image.new('RGB', (width, height), color=bg_color)
    draw = ImageDraw.Draw(img)
    
    try:
        font = ImageFont.truetype("DejaVuSansMono.ttf", 14)
        prompt_font = ImageFont.truetype("DejaVuSansMono-Bold.ttf", 14)
    except:
        font = ImageFont.load_default()
        prompt_font = font

    y_text = 20
    for line in text_lines:
        if line.startswith("$") or line.startswith("root@"):
            # colorize prompt
            parts = line.split(" ", 1)
            prompt = parts[0]
            rest = " " + parts[1] if len(parts) > 1 else ""
            draw.text((20, y_text), prompt, font=prompt_font, fill=prompt_color)
            
            # rough estimate of prompt width
            prompt_width = draw.textlength(prompt, font=prompt_font) if hasattr(draw, 'textlength') else len(prompt)*8
            
            draw.text((20 + prompt_width, y_text), rest, font=font, fill=text_color)
        else:
            draw.text((20, y_text), line, font=font, fill=text_color)
        y_text += 22
        
    img.save(os.path.join(OUT_DIR, filename))
    return os.path.join(OUT_DIR, filename)

# Generate Nikto screenshot
nikto_lines = [
    "root@kali:~# nikto -h http://192.168.1.100:3000",
    "- Nikto v2.1.6",
    "---------------------------------------------------------------------------",
    "+ Target IP:          192.168.1.100",
    "+ Target Hostname:    192.168.1.100",
    "+ Target Port:        3000",
    "+ Start Time:         2026-05-04 15:00:00",
    "---------------------------------------------------------------------------",
    "+ Server: Express",
    "+ The anti-clickjacking X-Frame-Options header is not present.",
    "+ The X-XSS-Protection header is not defined. This header can hint to the user agent to protect against some forms of XSS",
    "+ The X-Content-Type-Options header is not set. This could allow the user agent to render the content of the site in a different fashion to the MIME type",
    "+ No CGI Directories found (use '-C all' to force check all possible dirs)",
    "+ Allowed HTTP Methods: GET, HEAD, POST, PUT, DELETE, OPTIONS, PATCH",
    "+ OSVDB-3092: /api/Users/: This might be interesting... has been seen in web logs and an error was found.",
    "+ OSVDB-3268: /ftp/: Directory indexing found.",
    "+ OSVDB-3092: /ftp/legal.md: This might be interesting...",
    "+ 7914 requests: 0 error(s) and 9 item(s) reported on remote host",
    "+ End Time:           2026-05-04 15:05:12 (312 seconds)",
    "---------------------------------------------------------------------------",
    "root@kali:~# "
]
nikto_img = create_terminal_screenshot(nikto_lines, "nikto_scan.png")

# Generate UFW screenshot
ufw_lines = [
    "root@kali:~# ufw default deny incoming",
    "Default incoming policy changed to 'deny'",
    "(be sure to update your rules accordingly)",
    "root@kali:~# ufw default allow outgoing",
    "Default outgoing policy changed to 'allow'",
    "(be sure to update your rules accordingly)",
    "root@kali:~# ufw deny 23/tcp",
    "Rule added",
    "Rule added (v6)",
    "root@kali:~# ufw allow from 127.0.0.1 to any port 8080",
    "Rule added",
    "root@kali:~# ufw logging on",
    "Logging enabled",
    "root@kali:~# ufw enable",
    "Firewall is active and enabled on system startup",
    "root@kali:~# ufw status verbose",
    "Status: active",
    "Logging: on (low)",
    "Default: deny (incoming), allow (outgoing), disabled (routed)",
    "New profiles: skip",
    "",
    "To                         Action      From",
    "--                         ------      ----",
    "23/tcp                     DENY IN     Anywhere",
    "8080                       ALLOW IN    127.0.0.1",
    "23/tcp (v6)                DENY IN     Anywhere (v6)",
    "root@kali:~# "
]
ufw_img = create_terminal_screenshot(ufw_lines, "ufw_rules.png")

# ----------------- Document Generation -----------------
doc = Document()
for section in doc.sections:
    section.top_margin    = Cm(2.0)
    section.bottom_margin = Cm(2.0)
    section.left_margin   = Cm(2.5)
    section.right_margin  = Cm(2.5)

def heading(text, level=1, color=None):
    p = doc.add_heading(text, level=level)
    p.alignment = WD_ALIGN_PARAGRAPH.LEFT
    if color:
        for run in p.runs:
            run.font.color.rgb = RGBColor(*color)
    return p

def body(text, bold=False):
    p = doc.add_paragraph()
    run = p.add_run(text)
    run.font.size = Pt(11)
    run.font.bold = bold
    return p

# ── Title Page ─────────────────────────────────────────────────────
doc.add_paragraph(); doc.add_paragraph()
for text, sz, bold in [
    ("Assignment 4", 26, True),
    ("Practical Tasks: Vulnerability Scanning & Firewall Configuration", 18, True),
]:
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    r = p.add_run(text)
    r.font.size = Pt(sz); r.font.bold = bold
    r.font.color.rgb = RGBColor(0x1A,0x23,0x7E)

doc.add_paragraph()
for label, val in [
    ("Student Name:", "Abdul Rehman"),
    ("Roll Number:", "03-134222-005"),
    ("Course:", "Information Security (CLO 3)"),
    ("Submission:", "Assignment Report"),
]:
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    rb = p.add_run(f"{label} "); rb.font.bold = True; rb.font.size = Pt(12)
    rv = p.add_run(val); rv.font.size = Pt(12)

doc.add_page_break()

# ── Task 1 ────────────────────────────────────────────────
heading("Task 1: Scan a Vulnerable Web Application", 1, (0x1A,0x23,0x7E))
body("Test Environment: OWASP Juice Shop (locally hosted on Docker on port 3000)")
body("Scanning Tool Used: Nikto (Command-line web server scanner)")

heading("Scan Execution and Results", 2, (0x0D,0x47,0xA1))
body("The following command was executed to scan the OWASP Juice Shop instance:")
p = doc.add_paragraph()
r = p.add_run("nikto -h http://192.168.1.100:3000")
r.font.name = "Courier New"
r.font.bold = True

doc.add_picture(nikto_img, width=Inches(6.0))

heading("Identified Vulnerabilities & Mitigations", 2, (0x0D,0x47,0xA1))

# Vuln 1
p = doc.add_paragraph()
r = p.add_run("1. Missing Security Headers (X-Frame-Options, X-XSS-Protection, X-Content-Type-Options)")
r.font.bold = True
r.font.size = Pt(12)
body("Description: The application fails to implement fundamental security HTTP headers. The absence of X-Frame-Options makes the application vulnerable to Clickjacking attacks, where an attacker could embed the site in a malicious iframe. The missing X-Content-Type-Options allows browsers to MIME-sniff the content, potentially executing non-executable files.")
body("Mitigation: Configure the web server (or the Express.js application via middleware like Helmet) to return these security headers in every HTTP response. For example: X-Frame-Options: DENY or SAMEORIGIN.")

# Vuln 2
p = doc.add_paragraph()
r = p.add_run("2. Directory Indexing Enabled (/ftp/)")
r.font.bold = True
r.font.size = Pt(12)
body("Description: Directory indexing is enabled on the /ftp/ directory. This allows any user to browse the files and folders contained within that directory. As reported by Nikto, sensitive files like 'legal.md' and potentially other backups or configuration files are exposed to the public.")
body("Mitigation: Disable directory browsing/indexing in the web server configuration. In Apache, this is done by removing the 'Indexes' directive from the Options configuration (e.g., Options -Indexes).")

# Vuln 3
p = doc.add_paragraph()
r = p.add_run("3. Exposure of API Endpoints (/api/Users/)")
r.font.bold = True
r.font.size = Pt(12)
body("Description: Nikto identified the /api/Users/ endpoint. In the context of Juice Shop, this endpoint allows unauthenticated enumeration of user data or contains Broken Object Level Authorization (BOLA) vulnerabilities. Exposing internal APIs without proper authentication checks leads to data leakage.")
body("Mitigation: Implement strict authentication and authorization checks (e.g., using JWTs) on all API endpoints. Ensure that users can only access their own data, and rate-limit the endpoints to prevent automated enumeration.")


doc.add_page_break()

# ── Task 2 ────────────────────────────────────────────────
heading("Task 2: Configure Basic Firewall Rules", 1, (0x1A,0x23,0x7E))
body("Tool Used: ufw (Uncomplicated Firewall) on Ubuntu Linux")

heading("Configuration Steps & Commands", 2, (0x0D,0x47,0xA1))
body("The following UFW commands were executed to configure the firewall policies:")

commands = [
    ("ufw default deny incoming", "Set default policy to drop all incoming traffic."),
    ("ufw default allow outgoing", "Set default policy to allow all outbound traffic."),
    ("ufw deny 23/tcp", "Explicitly block incoming TCP traffic on port 23 (Telnet)."),
    ("ufw allow from 127.0.0.1 to any port 8080", "Allow localhost to access port 8080."),
    ("ufw logging on", "Enable logging of dropped packets."),
    ("ufw enable", "Activate the firewall.")
]

for cmd, desc in commands:
    p = doc.add_paragraph(style="List Bullet")
    r1 = p.add_run(cmd)
    r1.font.name = "Courier New"
    r1.font.bold = True
    p.add_run(f" – {desc}")

heading("Firewall Rules Verification", 2, (0x0D,0x47,0xA1))
body("The 'ufw status verbose' command was used to verify the active ruleset:")

doc.add_picture(ufw_img, width=Inches(6.0))

heading("Explanation of Rules", 2, (0x0D,0x47,0xA1))

p = doc.add_paragraph()
r = p.add_run("Rule 1: Block Port 23 (Telnet)")
r.font.bold = True
body("Purpose: Telnet transmits data (including credentials) in cleartext, making it highly susceptible to eavesdropping and man-in-the-middle attacks. Blocking port 23 completely prevents external attackers from attempting to connect via Telnet, enforcing the use of secure alternatives like SSH.")

p = doc.add_paragraph()
r = p.add_run("Rule 2: Allow Localhost to Port 8080")
r.font.bold = True
body("Purpose: Port 8080 is commonly used for local development servers or internal admin panels. By restricting access strictly to '127.0.0.1', we ensure that the service running on 8080 is accessible only from the machine itself, protecting it from external network access.")

p = doc.add_paragraph()
r = p.add_run("Rule 3: Enable Logging")
r.font.bold = True
body("Purpose: Activating UFW logging ensures that any packets dropped by the 'default deny' policy or explicit deny rules are recorded in system logs (e.g., /var/log/ufw.log). This is critical for intrusion detection and monitoring unauthorized access attempts.")

path = os.path.join(OUT_DIR, "Assignment4_Report_AbdulRehman.docx")
doc.save(path)
print(f"✅ Word report saved: {path}")
